import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_docx(markdown_file, output_file):
    """Convert Markdown file to DOCX with proper formatting"""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Heading 1 (##)
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True
            
        # Heading 2 (###)
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
            
        # Heading 3 (####)
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=3)
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.bold = True
        
        # Code block (```)
        elif line.startswith('```'):
            code_lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            
            # Add code block
            p = doc.add_paragraph()
            p.style = 'Normal'
            run = p.add_run('\\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.left_indent = Inches(0.5)
            
        # Horizontal rule (---)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Bullet list (-)
        elif line.startswith('- '):
            text = line[2:].strip()
            # Remove markdown formatting
            text = remove_markdown_formatting(text)
            doc.add_paragraph(text, style='List Bullet')
            
        # Bold header line (starts with **)
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
        
        # Regular paragraph
        else:
            # Check if it's a continuation of list
            if line.startswith('  - '):
                text = line[4:].strip()
                text = remove_markdown_formatting(text)
                doc.add_paragraph(text, style='List Bullet 2')
            else:
                text = remove_markdown_formatting(line)
                if text:
                    p = doc.add_paragraph(text)
                    p.style = 'Normal'
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"Document saved to: {output_file}")

def remove_markdown_formatting(text):
    """Remove markdown formatting from text"""
    # Remove bold
    text = re.sub(r'\\*\\*(.+?)\\*\\*', r'\\1', text)
    # Remove italic
    text = re.sub(r'\\*(.+?)\\*', r'\\1', text)
    # Remove code
    text = re.sub(r'`(.+?)`', r'\\1', text)
    # Remove links [text](url)
    text = re.sub(r'\\[(.+?)\\]\\(.+?\\)', r'\\1', text)
    return text

if __name__ == "__main__":
    markdown_file = "DOKUMENTASI_TEKNIS.md"
    output_file = "DOKUMENTASI_TEKNIS.docx"
    
    print("Converting Markdown to DOCX...")
    parse_markdown_to_docx(markdown_file, output_file)
    print("Conversion complete!")
